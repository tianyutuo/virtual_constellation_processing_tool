"""
分析查找表设置不同步长对插值结果的影响
输出到word表格
"""
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import re
import pickle
import math
import numpy as np
from glob import iglob
from create_LUT import permutate_invars


def sixs_step_analysis(iLUT_fpath,
                       docx_file,
                       analysis_para,
                       start, end, r,
                       step_sets,
                       analysis_step,
                       invars):
    os.chdir(iLUT_fpath)
    # ---------------------------------------------------------------------------------------------------
    # 搜寻分析参数不同步长设置的插值查找表
    # ---------------------------------------------------------------------------------------------------
    blue_fpaths = sorted(list(iglob(os.path.join(iLUT_fpath, '**',
                                                 'wavelength_0.45_0.52_f_' + analysis_para + '_step*.ilut'),
                                    recursive=True)))
    green_fpaths = sorted(list(iglob(os.path.join(iLUT_fpath, '**',
                                                  'wavelength_0.52_0.59_f_' + analysis_para + '_step*.ilut'),
                                     recursive=True)))
    red_fpaths = sorted(list(iglob(os.path.join(iLUT_fpath, '**',
                                                'wavelength_0.63_0.69_f_' + analysis_para + '_step*.ilut'),
                                   recursive=True)))
    nir_fpaths = sorted(list(iglob(os.path.join(iLUT_fpath, '**',
                                                'wavelength_0.77_0.89_f_' + analysis_para + '_step*.ilut'),
                                   recursive=True)))
    # 读取插值查找表
    step_num = len(step_sets)
    blue_iLUTs = {}
    green_iLUTs = {}
    red_iLUTs = {}
    nir_iLUTs = {}
    for i in range(step_num):
        with open(blue_fpaths[i], 'rb') as blue_ilut_file:
            blue_iLUTs[re.findall(r'step\d+', blue_fpaths[i])[0]] = pickle.load(blue_ilut_file)
        with open(green_fpaths[i], 'rb') as green_ilut_file:
            green_iLUTs[re.findall(r'step\d+', green_fpaths[i])[0]] = pickle.load(green_ilut_file)
        with open(red_fpaths[i], 'rb') as red_ilut_file:
            red_iLUTs[re.findall(r'step\d+', red_fpaths[i])[0]] = pickle.load(red_ilut_file)
        with open(nir_fpaths[i], 'rb') as nir_ilut_file:
            nir_iLUTs[re.findall(r'step\d+', red_fpaths[i])[0]] = pickle.load(nir_ilut_file)

    # ---------------------------------------------------------------------------------------------------
    # 查找表插值验证：6S模型直接输出的结果
    # ---------------------------------------------------------------------------------------------------
    blue_validate_file = 'validate/wavelength_0.45_0.52_f_' + analysis_para + '_validate.lut'
    green_validate_file = 'validate/wavelength_0.52_0.59_f_' + analysis_para + '_validate.lut'
    red_validate_file = 'validate/wavelength_0.63_0.69_f_' + analysis_para + '_validate.lut'
    nir_validate_file = 'validate/wavelength_0.77_0.89_f_' + analysis_para + '_validate.lut'
    # 读取验证数据
    blue_validates = pickle.load(open(blue_validate_file, 'rb'))['outputs']
    green_validates = pickle.load(open(green_validate_file, 'rb'))['outputs']
    red_validates = pickle.load(open(red_validate_file, 'rb'))['outputs']
    nir_validates = pickle.load(open(nir_validate_file, 'rb'))['outputs']

    # ---------------------------------------------------------------------------------------------------
    # 在word文档中建立表格
    # ---------------------------------------------------------------------------------------------------
    if not os.path.exists(docx_file):
        d = Document()
    else:
        d = Document(docx_file)
    t = d.add_table(rows=math.ceil((end - start) / r) * step_num + 2, cols=6, style='Table Grid')
    # 设置每个单元格居中显示
    for row in range(math.ceil((end - start) / r) * step_num + 2):
        for col in range(6):
            t.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER     # 垂直居中
            t.cell(row, col).paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # 水平居中
    param_cell1 = t.cell(0, 0)
    param_cell1.text = analysis_para
    param_cell2 = t.cell(1, 0)
    param_cell1.merge(param_cell2)
    step_cell1 = t.cell(0, 1)
    step_cell1.text = '步长'
    step_cell2 = t.cell(1, 1)
    step_cell1.merge(step_cell2)
    relative_error_cell1 = t.cell(0, 2)
    relative_error_cell1.text = '平均相对误差（Mean Relative Error）'
    relative_error_cell2 = t.cell(0, 5)
    relative_error_cell1.merge(relative_error_cell2)
    blue_cell = t.cell(1, 2)
    blue_cell.text = 'Blue'
    green_cell = t.cell(1, 3)
    green_cell.text = 'Green'
    red_cell = t.cell(1, 4)
    red_cell.text = 'Red'
    nir_cell = t.cell(1, 5)
    nir_cell.text = 'NIR'
    # ---------------------------------------------------------------------------------------------------
    # 计算相对误差
    # ---------------------------------------------------------------------------------------------------
    relative_error = {}
    for step_set in step_sets:
        relative_error['blue_' + step_set] = []
        relative_error['green_' + step_set] = []
        relative_error['red_' + step_set] = []
        relative_error['nir_' + step_set] = []
    perms = permutate_invars(invars)
    for i in range(int((end - start) / analysis_step) + 1):
        blue_validate = blue_validates[i]
        green_validate = green_validates[i]
        red_validate = red_validates[i]
        nir_validate = nir_validates[i]

        perm = perms[i]
        for step_set in step_sets:
            blue_interpolated = blue_iLUTs[step_set](perm[0], perm[1], perm[2], perm[3], perm[4], perm[5], perm[6])
            green_interpolated = green_iLUTs[step_set](perm[0], perm[1], perm[2], perm[3], perm[4], perm[5], perm[6])
            red_interpolated = red_iLUTs[step_set](perm[0], perm[1], perm[2], perm[3], perm[4], perm[5], perm[6])
            nir_interpolated = nir_iLUTs[step_set](perm[0], perm[1], perm[2], perm[3], perm[4], perm[5], perm[6])
            relative_error['blue_' + step_set].append((blue_validate - blue_interpolated) / blue_validate)
            relative_error['green_' + step_set].append((green_validate - green_interpolated) / green_validate)
            relative_error['red_' + step_set].append((red_validate - red_interpolated) / red_validate)
            relative_error['nir_' + step_set].append((nir_validate - nir_interpolated) / nir_validate)
    # ---------------------------------------------------------------------------------------------------
    # 将分析结果写入表格
    # ---------------------------------------------------------------------------------------------------
    for i in range(math.ceil((end - start) / r)):
        range_cell1 = t.cell(i * step_num + 2, 0)
        if (i + 1) * r < end:
            range_cell1.text = f'{i * r}-{(i + 1) * r}'
        else:
            range_cell1.text = f'{i * r}-{end}'
        range_cell2 = t.cell((i + 1) * step_num + 1, 0)
        range_cell1.merge(range_cell2)
        for j in range(step_num):
            r_num = int(r / analysis_step)

            step_set_cell = t.cell(i * step_num + 2 + j, 1)
            step_set_cell.text = step_sets[j]
            blue_relative_error_cell = t.cell(i * step_num + 2 + j, 2)
            green_relative_error_cell = t.cell(i * step_num + 2 + j, 3)
            red_relative_error_cell = t.cell(i * step_num + 2 + j, 4)
            nir_relative_error_cell = t.cell(i * step_num + 2 + j, 5)
            if (i + 1) * r < end:
                blue_relative_error = np.sum(relative_error['blue_' + step_sets[j]][i * r_num:(i + 1) * r_num],
                                             axis=0)
                green_relative_error = np.sum(relative_error['green_' + step_sets[j]][i * r_num:(i + 1) * r_num],
                                              axis=0)
                red_relative_error = np.sum(relative_error['red_' + step_sets[j]][i * r_num:(i + 1) * r_num],
                                            axis=0)
                nir_relative_error = np.sum(relative_error['nir_' + step_sets[j]][i * r_num:(i + 1) * r_num],
                                            axis=0)
            else:
                blue_relative_error = np.sum(relative_error['blue_' + step_sets[j]][i * r_num:], axis=0)
                green_relative_error = np.sum(relative_error['green_' + step_sets[j]][i * r_num:], axis=0)
                red_relative_error = np.sum(relative_error['red_' + step_sets[j]][i * r_num:], axis=0)
                nir_relative_error = np.sum(relative_error['nir_' + step_sets[j]][i * r_num:], axis=0)
                r_num = int((end - i * r) / analysis_step) + 1
            blue_relative_error_cell.text = 'xa: %.2f%%\nxb: %.2f%%\nxc: %.2f%%' % \
                                            (blue_relative_error[0] / r_num * 100,
                                             blue_relative_error[1] / r_num * 100,
                                             blue_relative_error[2] / r_num * 100)
            green_relative_error_cell.text = 'xa: %.2f%%\nxb: %.2f%%\nxc: %.2f%%' % \
                                             (green_relative_error[0] / r_num * 100,
                                              green_relative_error[1] / r_num * 100,
                                              green_relative_error[2] / r_num * 100)
            red_relative_error_cell.text = 'xa: %.2f%%\nxb: %.2f%%\nxc: %.2f%%' % \
                                           (red_relative_error[0] / r_num * 100,
                                            red_relative_error[1] / r_num * 100,
                                            red_relative_error[2] / r_num * 100)
            nir_relative_error_cell.text = 'xa: %.2f%%\nxb: %.2f%%\nxc: %.2f%%' % \
                                           (nir_relative_error[0] / r_num * 100,
                                            nir_relative_error[1] / r_num * 100,
                                            nir_relative_error[2] / r_num * 100)

    d.add_paragraph('')   # 间隔
    d.save(docx_file)


if __name__ == '__main__':
    # 水汽步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='h2o',
                       start=0, end=8.5, r=1,
                       step_sets=['step025', 'step05', 'step10', 'step20', 'step35'],
                       analysis_step=0.1,
                       invars={
                           'sza': [50],
                           'vza': [48],
                           'raa': [48],
                           'aod': [0.3],
                           'h2o': np.arange(0, 8.6, 0.1),
                           'o3': [0.4],
                           'alt': [1]
                       })

    # 臭氧步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='o3',
                       start=0, end=0.8, r=0.1,
                       step_sets=['step005', 'step01', 'step02', 'step04', 'step08'],
                       analysis_step=0.01,
                       invars={
                           'sza': [50],
                           'vza': [48],
                           'raa': [48],
                           'aod': [0.3],
                           'h2o': [1],
                           'o3': np.arange(0, 0.81, 0.01),
                           'alt': [1]
                       })

    # SZA步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='sza_angle',
                       start=0, end=80, r=10,
                       step_sets=['step2', 'step5', 'step10'],
                       analysis_step=1,
                       invars={
                           'sza': np.arange(0, 81, 1),
                           'vza': [48],
                           'raa': [48],
                           'aod': [0.3],
                           'h2o': [1],
                           'o3': [0.4],
                           'alt': [1]
                       })

    # VZA步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='vza_angle',
                       start=0, end=80, r=10,
                       step_sets=['step2', 'step5', 'step10'],
                       analysis_step=1,
                       invars={
                           'sza': [50],
                           'vza': np.arange(0, 81, 1),
                           'raa': [48],
                           'aod': [0.3],
                           'h2o': [1],
                           'o3': [0.4],
                           'alt': [1]
                       })

    # RAA步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='raa_angle',
                       start=0, end=180, r=10,
                       step_sets=['step10', 'step20', 'step30'],
                       analysis_step=1,
                       invars={
                           'sza': [50],
                           'vza': [48],
                           'raa': np.arange(0, 181, 1),
                           'aod': [0.3],
                           'h2o': [1],
                           'o3': [0.4],
                           'alt': [1]
                       })

    # AOD步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='aod',
                       start=0, end=3.0, r=0.1,
                       step_sets=['step005', 'step01', 'step025', 'step05', 'step075', 'step1'],
                       analysis_step=0.01,
                       invars={
                           'sza': [50],
                           'vza': [48],
                           'raa': [48],
                           'aod': np.arange(0, 3.01, 0.01),
                           'h2o': [1],
                           'o3': [0.4],
                           'alt': [1]
                       })

    # 目标高度步长设置分析
    sixs_step_analysis(iLUT_fpath=r'F:\Experiment\6S_parameters_step_analysis\step_analysis',
                       docx_file=r'D:\虚拟星座\高分数据\大气校正实验\6S查找表步长分析表格_test.docx',
                       analysis_para='alt',
                       start=0, end=8.0, r=1.0,
                       step_sets=['step1', 'step2', 'step3', 'step4'],
                       analysis_step=0.1,
                       invars={
                           'sza': [50],
                           'vza': [48],
                           'raa': [48],
                           'aod': [0.3],
                           'h2o': [1],
                           'o3': [0.4],
                           'alt': np.arange(0, 8.1, 0.1)
                       })
